import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Function to add borders to the table in the Word document
def add_borders_to_table(table):
    for row in table.rows:
        for cell in row.cells:
            cell.border = {
                'top': {'sz': 4, 'space': 0, 'color': '000000', 'val': 'single'},
                'bottom': {'sz': 4, 'space': 0, 'color': '000000', 'val': 'single'},
                'left': {'sz': 4, 'space': 0, 'color': '000000', 'val': 'single'},
                'right': {'sz': 4, 'space': 0, 'color': '000000', 'val': 'single'},
            }

# App title
st.title("Dummy Table with Categorical and Numeric Variables")

# Sidebar for navigation
st.sidebar.title("Navigation")
section = st.sidebar.radio("Go to", ["Test Overview", "Test Illustration"])

# 1. Test Overview Section
if section == "Test Overview":
    st.header("Overview: Dummy Table")

    st.subheader("What is a Dummy Table?")
    st.write("""
        A dummy table summarizes data involving both categorical and numeric variables.
        It provides insights into how numeric values interact with categorical groups.
    """)

# 2. Test Illustration Section
elif section == "Test Illustration":
    st.header("Illustration: Dummy Table")
    
    st.subheader("Upload your dataset (CSV or XLSX)")
    uploaded_file = st.file_uploader("Choose a CSV or XLSX file", type=["csv", "xlsx"])
    
    if uploaded_file is not None:
        # Load the dataset based on file type
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            elif uploaded_file.name.endswith('.xlsx'):
                df = pd.read_excel(uploaded_file)
                
            st.write("Here is a preview of your data:")
            st.write(df.head())

            # Ask the user to select one categorical and multiple numeric columns
            cat_column = st.selectbox("Select a categorical variable", df.select_dtypes(include=['object', 'category']).columns)
            num_columns = st.multiselect("Select numeric variables", df.select_dtypes(include=['float64', 'int64']).columns)

            # Input field for the title of the dummy tables
            table_title = st.text_input("Enter a title for the Dummy Tables:", "Summary Table")

            # Create a Word document to save outputs
            doc = Document()

            for num_col in num_columns:
                # Generate the summary table for the current numeric column
                summary_table = df.groupby(cat_column)[num_col].agg(['count', 'mean', 'sum', 'std']).reset_index()
                summary_table.columns = [cat_column, 'Count', 'Mean', 'Total', 'Std Dev']

                # Format Mean and Std Dev to one decimal place
                summary_table['Mean'] = summary_table['Mean'].round(1)
                summary_table['Std Dev'] = summary_table['Std Dev'].round(1)

                # Calculate percentage
                summary_table['Percentage'] = (summary_table['Total'] / summary_table['Total'].sum() * 100).round(1)

                # Display the summary table
                st.write(f"**{table_title} for {num_col}**")
                st.write(summary_table)

                # Input fields for plot titles
                bar_chart_title = st.text_input(f"Enter a title for the Horizontal Bar Chart for {num_col}:", f"Mean of {num_col} by {cat_column}")
                percentage_chart_title = st.text_input(f"Enter a title for the Horizontal Bar Chart for {num_col}:", f"Percentage of {num_col} by {cat_column}")
                pie_chart_title = st.text_input(f"Enter a title for the Pie Chart for {num_col}:", f"Percentage of {num_col} by {cat_column}")

                # Horizontal Bar Chart of Mean
                st.write(f"Horizontal Bar Chart of Mean for {num_col} by {cat_column}:")
                mean_values = summary_table.sort_values(by='Mean', ascending=False)
                plt.figure(figsize=(10, 6))
                plt.barh(mean_values[cat_column], mean_values['Mean'], color='skyblue')
                plt.xlabel('Mean')
                plt.title(bar_chart_title)
                plt.grid(axis='x')
                plt.tight_layout()
                plt.savefig(f'mean_bar_chart_{num_col}.png')
                st.pyplot(plt)


                # Horizontal Bar Chart of Percentage
                st.write(f"Percentage for {num_col} by {cat_column}:")
                percentage_values = summary_table.sort_values(by='Percentage', ascending=False)
                plt.figure(figsize=(10, 6))
                plt.barh(percentage_values[cat_column], percentage_values['Percentage'], color='skyblue')
                plt.xlabel('Percentage')
                plt.title(percentage_chart_title)
                plt.grid(axis='x')
                plt.tight_layout()
                plt.savefig(f'mean_bar_chart_{num_col}.png')
                st.pyplot(plt)

                

                # Pie Chart of Total with rotated labels
                st.write(f"Percentage of {num_col} by {cat_column}:")
                plt.figure(figsize=(8, 8))
                wedges, texts, autotexts = plt.pie(mean_values['Total'], labels=mean_values[cat_column], autopct='%1.1f%%', startangle=90)
                
                # Rotate labels
                for text in texts:
                    text.set_rotation(45)

                plt.title(pie_chart_title)
                plt.tight_layout()
                plt.savefig(f'total_pie_chart_{num_col}.png')
                st.pyplot(plt)

                # Save the dummy table as a picture
                fig, ax = plt.subplots(figsize=(8, len(summary_table) * 0.5))  # Adjust size based on number of rows
                ax.axis('tight')
                ax.axis('off')
                the_table = ax.table(cellText=summary_table.values, colLabels=summary_table.columns, cellLoc='center', loc='center')
                the_table.auto_set_font_size(False)
                the_table.set_fontsize(10)
                the_table.scale(1.2, 1.2)

                # Save the table as a PNG image
                plt.title(f'{table_title} - {num_col}')
                plt.savefig(f'dummy_table_{num_col}.png', bbox_inches='tight', pad_inches=0.1)
                plt.close(fig)  # Close the figure to avoid display

                # Add to Word document with tabular format
                doc.add_heading(f'{table_title} - {num_col}', level=2)
                
                # Create a table in the Word document
                word_table = doc.add_table(rows=1, cols=len(summary_table.columns))
                hdr_cells = word_table.rows[0].cells
                for i, column_name in enumerate(summary_table.columns):
                    hdr_cells[i].text = column_name
                
                # Fill in the table with data and add borders
                for _, row in summary_table.iterrows():
                    row_cells = word_table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # Add borders to the table
                add_borders_to_table(word_table)

                # Add plots to the Word document
                doc.add_heading(f'Mean Bar Chart - {num_col}', level=3)
                doc.add_picture(f'mean_bar_chart_{num_col}.png', width=Inches(5))
                
                doc.add_heading(f'Total Pie Chart - {num_col}', level=3)
                doc.add_picture(f'total_pie_chart_{num_col}.png', width=Inches(5))

                # Add table image to the Word document
                doc.add_heading(f'Dummy Table Image - {num_col}', level=3)
                doc.add_picture(f'dummy_table_{num_col}.png', width=Inches(5))

            # Save the document
            doc.save('dummy_tables.docx')

            # Provide download link for the Word document
            with open('dummy_tables.docx', 'rb') as f:
                st.download_button(
                    label="Download Dummy Tables as Word Document",
                    data=f,
                    file_name='dummy_tables.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

            st.success("Plots and tables have been generated and saved as images and the Word document has been created.")

        except Exception as e:
            st.error(f"Error loading file: {e}")
