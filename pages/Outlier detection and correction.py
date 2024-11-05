import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import zipfile
from io import BytesIO
import re
from docx import Document
from docx.shared import Inches

# Sidebar for Data Management
st.sidebar.title("Data Management")
data_management_option = st.sidebar.radio(
    "Choose an option:",
    ["Outlier Detection and Correction"]
)

# Initialize session state variables
if 'df' not in st.session_state:
    st.session_state.df = None
if 'filtered_df' not in st.session_state:
    st.session_state.filtered_df = None

# Function to filter dataframe based on categorical columns and their unique values
def filter_dataframe(df, categorical_columns, values):
    condition = pd.Series([True] * len(df))  # Initialize a condition with True values
    for col, val in zip(categorical_columns, values):
        condition &= (df[col].astype(str) == str(val))  # Convert to string for comparison
    return df[condition]

# Function to detect outliers using Scatterplot with Q1 and Q3 lines
def detect_outliers_scatterplot(df, col):
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return df[(df[col] < lower_bound) | (df[col] > upper_bound)], lower_bound, upper_bound

# Outlier correction methods
def replace_outliers_with_mean(df, col):
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    mean_value = df[col].mean()
    df[col] = df[col].where((df[col] >= lower_bound) & (df[col] <= upper_bound), mean_value)
    return df

def replace_outliers_with_mean_excluding(df, col):
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    mean_value = df[(df[col] >= lower_bound) & (df[col] <= upper_bound)][col].mean()
    df[col] = df[col].where((df[col] >= lower_bound) & (df[col] <= upper_bound), mean_value)
    return df

def replace_outliers_with_median(df, col):
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    median_value = df[col].median()
    df[col] = df[col].where((df[col] >= lower_bound) & (df[col] <= upper_bound), median_value)
    return df

def replace_outliers_with_median_excluding(df, col):
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    median_value = df[(df[col] >= lower_bound) & (df[col] <= upper_bound)][col].median()
    df[col] = df[col].where((df[col] >= lower_bound) & (df[col] <= upper_bound), median_value)
    return df

def winsorize_column(df, col):
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    df[col] = df[col].clip(lower=lower_bound, upper=upper_bound)
    return df

# Function to calculate moving average
def calculate_moving_avg(series, window):
    return series.rolling(window=window, min_periods=1).mean()

def calculate_moving_avg_excluding_outliers(series, window):
    Q1 = series.quantile(0.25)
    Q3 = series.quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return series.where((series >= lower_bound) & (series <= upper_bound)).rolling(window=window, min_periods=1).mean()

# Function to create a subplot for outlier detection and correction methods
def create_subplot_for_outlier_detection_and_correction(filtered_df, numeric_column, adm1, adm3, Year, hf):
    outliers, lower_bound, upper_bound = detect_outliers_scatterplot(filtered_df, numeric_column)

    # Create subplots (4 rows, 2 columns)
    fig, axs = plt.subplots(4, 2, figsize=(15, 15))
    axs = axs.flatten()

    # Scatter plot for original data with outlier bounds
    axs[0].scatter(filtered_df['year_mon'], filtered_df[numeric_column], color='blue', label='Data Points')
    if not outliers.empty:
        axs[0].scatter(outliers['year_mon'], outliers[numeric_column], color='red', label='Outliers')
    axs[0].axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
    axs[0].axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
    axs[0].set_title(f'Outlier Detection for {numeric_column}')

    # Replace outliers with mean including outliers
    mean_incld_df = replace_outliers_with_mean(filtered_df.copy(), numeric_column)
    axs[1].scatter(mean_incld_df['year_mon'], mean_incld_df[numeric_column], color='blue', label='Corrected Data')
    axs[1].axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
    axs[1].axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
    axs[1].set_title(f'Mean (Including Outliers) for {numeric_column}')

    # Replace outliers with mean excluding outliers
    mean_excld_df = replace_outliers_with_mean_excluding(filtered_df.copy(), numeric_column)
    axs[2].scatter(mean_excld_df['year_mon'], mean_excld_df[numeric_column], color='blue', label='Corrected Data')
    axs[2].axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
    axs[2].axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
    axs[2].set_title(f'Mean (Excluding Outliers) for {numeric_column}')

    # Replace outliers with median including outliers
    median_incld_df = replace_outliers_with_median(filtered_df.copy(), numeric_column)
    axs[3].scatter(median_incld_df['year_mon'], median_incld_df[numeric_column], color='blue', label='Corrected Data')
    axs[3].axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
    axs[3].axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
    axs[3].set_title(f'Median (Including Outliers) for {numeric_column}')

    # Replace outliers with median excluding outliers
    median_excld_df = replace_outliers_with_median_excluding(filtered_df.copy(), numeric_column)
    axs[4].scatter(median_excld_df['year_mon'], median_excld_df[numeric_column], color='blue', label='Corrected Data')
    axs[4].axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
    axs[4].axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
    axs[4].set_title(f'Median (Excluding Outliers) for {numeric_column}')

    # Winsorization
    winsorized_df = winsorize_column(filtered_df.copy(), numeric_column)
    axs[5].scatter(winsorized_df['year_mon'], winsorized_df[numeric_column], color='blue', label='Winsorized Data')
    axs[5].axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
    axs[5].axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
    axs[5].set_title(f'Winsorization for {numeric_column}')

    # Moving average including all values
    moving_avg_incld = calculate_moving_avg(filtered_df[numeric_column], window=3)
    axs[6].scatter(filtered_df['year_mon'], moving_avg_incld, color='blue', label='Moving Average (Included)')
    axs[6].axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
    axs[6].axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
    axs[6].set_title(f'Moving Average (Including Outliers) for {numeric_column}')

    # Moving average excluding outliers
    moving_avg_excld = calculate_moving_avg_excluding_outliers(filtered_df[numeric_column], window=3)
    axs[7].scatter(filtered_df['year_mon'], moving_avg_excld, color='blue', label='Moving Average (Excluded)')
    axs[7].axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
    axs[7].axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
    axs[7].set_title(f'Moving Average (Excluding Outliers) for {numeric_column}')

    # Set labels for all subplots
    for ax in axs:
        ax.set_xlabel('Year-Month')
        ax.set_ylabel(numeric_column)

    # Set legend for the entire figure
    handles, labels = axs[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc='upper left', bbox_to_anchor=(-0.1, 1.05), borderaxespad=0.)

    plt.tight_layout()
    fig.suptitle(f'Outlier Detection and Correction for adm1 = {adm1}, adm3 = {adm3}, Year = {Year}, hf = {hf}, Variable = {numeric_column}', y=1.02)
    return fig

# Main App
st.title("Outlier Detection and Correction")

if data_management_option == "Outlier Detection and Correction":
    st.header("Select Columns for Outlier Detection")

    if st.session_state.df is None:
        st.warning("No dataset available. Please upload your dataset.")
        uploaded_file = st.file_uploader("Upload a CSV or Excel file", type=["csv", "xlsx"])
        if uploaded_file:
            if uploaded_file.name.endswith(".csv"):
                st.session_state.df = pd.read_csv(uploaded_file)
            else:
                st.session_state.df = pd.read_excel(uploaded_file)
            st.session_state.df['year_mon'] = pd.to_datetime(st.session_state.df['Year'].astype(str) + '-' + st.session_state.df['Month'].astype(str), format='%Y-%m')
            st.success("Dataset uploaded successfully!")
            st.dataframe(st.session_state.df)

    if st.session_state.df is not None:
        # User selects adm1, adm3, Year, Month, hf
        adm1 = st.selectbox("Select adm1:", st.session_state.df['adm1'].unique())
        adm3_options = st.session_state.df[st.session_state.df['adm1'] == adm1]['adm3'].unique()
        adm3 = st.selectbox("Select adm3:", adm3_options)
        Year_options = st.session_state.df[(st.session_state.df['adm1'] == adm1) & (st.session_state.df['adm3'] == adm3)]['Year'].unique()
        Year = st.selectbox("Select Year:", Year_options)
        Month_options = st.session_state.df[(st.session_state.df['adm1'] == adm1) & (st.session_state.df['adm3'] == adm3) & (st.session_state.df['Year'] == Year)]['Month'].unique()
        hf_options = st.session_state.df[(st.session_state.df['adm1'] == adm1) & (st.session_state.df['adm3'] == adm3) & (st.session_state.df['Year'] == Year) & (st.session_state.df['Month'].isin(Month_options))]['hf'].unique()
        hf = st.multiselect("Select hf:", hf_options)

        # Filter data based on selected values
        filtered_df = st.session_state.df[(st.session_state.df['adm1'] == adm1) &
                                          (st.session_state.df['adm3'] == adm3) &
                                          (st.session_state.df['Year'] == Year) &
                                          (st.session_state.df['hf'].isin(hf))]

        # User selects the numeric column for outlier detection and correction
        numeric_column = st.selectbox("Select a Numeric Column for Outlier Detection:", filtered_df.select_dtypes(include='number').columns)

        if numeric_column:
            st.write("Outliers will be detected and visualized for the selected categorical values.")

            if st.button("Generate Outlier Report"):
                word_buffer = BytesIO()  # Create a buffer for the Word file
                excel_buffer = BytesIO()  # Create a buffer for the Excel file
                document = Document()

                outlier_data = pd.ExcelWriter(excel_buffer, engine='xlsxwriter')

                # Prepare the zip archive for the output files
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w') as zf:
                    # Create a subplot for detection and correction for each unique hf
                    unique_hfs = hf
                    for unique_hf in unique_hfs:
                        hf_filtered_df = filtered_df[filtered_df['hf'] == unique_hf]
                        fig = create_subplot_for_outlier_detection_and_correction(hf_filtered_df, numeric_column, adm1, adm3, Year, unique_hf)

                        # Display the figure in Streamlit
                        st.pyplot(fig)

                        # Save plot as image to add to Word document
                        img_buffer = BytesIO()
                        fig.savefig(img_buffer, format='png')
                        img_buffer.seek(0)

                        # Add a paragraph before the image
                        document.add_paragraph(f'Outlier Detection and Correction for adm1 = {adm1}, adm3 = {adm3}, Year = {Year}, hf = {unique_hf}')
                        document.add_picture(img_buffer, width=Inches(5))
                        plt.close(fig)  # Close the plot to avoid memory issues

                    # Save the Word document
                    document.save(word_buffer)

                    # Close the Excel writer (finalize the file)
                    outlier_data.close()

                    # Write Word and Excel to zip
                    zf.writestr('outliers_report.docx', word_buffer.getvalue())
                    zf.writestr('outliers_data.xlsx', excel_buffer.getvalue())

                st.success("Outliers processed successfully!")
                # Offer the zip file for download
                st.download_button(
                    label="Download ZIP file with Word Report and Outliers",
                    data=zip_buffer.getvalue(),
                    file_name='outliers_results.zip',
                    mime='application/zip'
                )
