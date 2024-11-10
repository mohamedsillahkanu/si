import streamlit as st
import seaborn as sns
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Inches  # Corrected import for Word export

# Function to create a subplot based on user-selected parameters
def create_subplot(ax, data, plot_type, feature1, feature2=None, title='', xlabel='', ylabel=''):
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    if plot_type == 'Bar Chart':
        if data[feature1].dtype == 'object':  # Ensure it's a categorical variable
            count_data = data[feature1].value_counts()
            sns.barplot(x=count_data.index, y=count_data.values, ax=ax)
            ax.set_title(title if title else 'Bar Chart')
        else:
            st.error(f"'{feature1}' is not a categorical variable. Please select a categorical variable for the bar chart.")
    elif plot_type == 'Pie Chart':
        pie_data = data[feature1].value_counts()
        ax.pie(pie_data, labels=pie_data.index, autopct='%1.1f%%')
        ax.set_title(title if title else 'Pie Chart')
    elif plot_type == 'Histogram':
        sns.histplot(data[feature1], kde=True, ax=ax)
        ax.set_title(title if title else 'Histogram')
    elif plot_type == 'Violin Plot':
        sns.violinplot(x=data[feature1], ax=ax)
        ax.set_title(title if title else 'Violin Plot')
    elif plot_type == 'Line Plot':
        ax.plot(data[feature1])
        ax.set_title(title if title else 'Line Plot')
    elif plot_type == 'Hexbin Plot' and feature2 is not None:
        ax.hexbin(data[feature1], data[feature2], gridsize=20, cmap='Blues')
        ax.set_title(title if title else 'Hexbin Plot')
    elif plot_type == 'Box Plot':
        sns.boxplot(x=data[feature1], ax=ax)
        ax.set_title(title if title else 'Box Plot')
    elif plot_type == 'Scatter Plot' and feature2 is not None:
        ax.scatter(data[feature1], data[feature2])
        ax.plot([data[feature1].min(), data[feature1].max()], [data[feature1].min(), data[feature1].max()], 'r--')
        ax.set_title(title if title else 'Scatter Plot')

# Function to generate the subplots based on user inputs
def generate_subplots(rows, cols, data, plot_types, features1, features2, titles, xlabels, ylabels):
    fig, axes = plt.subplots(rows, cols, figsize=(5*cols, 4*rows))

    # If only one subplot, make axes a list
    if rows == 1 and cols == 1:
        axes = [axes]
    elif rows == 1 or cols == 1:
        axes = axes.flatten()  # For single row or single column layouts
    else:
        axes = axes.flatten()  # Flatten axes array for easy iteration

    for i in range(rows * cols):
        if i < len(plot_types):
            plot_type = plot_types[i]
            feature1 = features1[i]
            feature2 = features2[i] if len(features2) > i else None
            title = titles[i] if len(titles) > i else ''
            xlabel = xlabels[i] if len(xlabels) > i else ''
            ylabel = ylabels[i] if len(ylabels) > i else ''
            create_subplot(axes[i], data, plot_type, feature1, feature2, title, xlabel, ylabel)
        else:
            axes[i].axis('off')  # Turn off any unused subplots

    plt.subplots_adjust(wspace=0.4, hspace=0.4)
    st.pyplot(fig)

    return fig  # Return the figure object to export later

# Function to export the dashboard to a Word document
def export_to_word(figures, filename="dashboard.docx"):
    doc = Document()
    doc.add_heading('Dashboard Export', 0)

    for i, fig in enumerate(figures):
        img_stream = BytesIO()
        fig.savefig(img_stream, format='png')
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(6))  # Use Inches from docx.shared
        doc.add_paragraph(f'Figure {i+1}')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(label="Download as Word", data=buffer, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Function to export the dashboard to a PDF document
def export_to_pdf(figures, filename="dashboard.pdf"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Dashboard Export", ln=True, align='C')

    for i, fig in enumerate(figures):
        img_stream = BytesIO()
        fig.savefig(img_stream, format='png')
        img_stream.seek(0)
        pdf.add_page()
        pdf.image(img_stream, x=10, y=20, w=180)
        pdf.ln(85)  # Adjust this value to control spacing between images
        pdf.cell(0, 10, f"Figure {i+1}", ln=True)

    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    st.download_button(label="Download as PDF", data=buffer, file_name=filename, mime="application/pdf")

# Streamlit app layout
# Initialize df as an empty DataFrame
df = pd.DataFrame()

# Streamlit app layout
st.title("Customizable Data Visualization Dashboard with Export Option")

# Sidebar for data source selection
st.sidebar.header("Data Source")
data_source = st.sidebar.radio("Choose data source", ["Use Sample Data", "Upload Excel/CSV"])

# Sidebar for number of pages, rows, and columns
n_pages = st.sidebar.number_input("Select number of pages", min_value=1, max_value=5, value=1)
n_rows = st.sidebar.number_input("Select number of rows", min_value=1, max_value=5, value=1)
n_cols = st.sidebar.number_input("Select number of columns", min_value=1, max_value=5, value=1)

if data_source == "Upload Excel/CSV":
    uploaded_file = st.sidebar.file_uploader("Upload an Excel or CSV file", type=["xlsx", "csv"])
    if uploaded_file:
        if uploaded_file.name.endswith(".xlsx"):
            df = pd.read_excel(uploaded_file)
        else:
            df = pd.read_csv(uploaded_file)
        st.write(df.head())  # Display the first few rows to confirm upload
else:
    # Sample data
    df = pd.DataFrame({
        'Feature 1': np.random.choice(['A', 'B', 'C', 'D'], 100),
        'Feature 2': np.random.randn(100),
        'Feature 3': np.random.randint(1, 10, 100),
        'Feature 4': np.random.randn(100),
    })

# User input for selecting features and plot types for each page
page_data = []
for page in range(n_pages):
    st.sidebar.subheader(f"Page {page+1} Configuration")
    plot_types = []
    features1 = []
    features2 = []
    titles = []
    xlabels = []
    ylabels = []

    for i in range(n_rows * n_cols):
        st.sidebar.subheader(f"Subplot {i+1} on Page {page+1}")
        plot_type = st.sidebar.selectbox(f"Select plot type for subplot {i+1} on Page {page+1}",
                                         ['Bar Chart', 'Pie Chart', 'Histogram', 'Violin Plot', 'Line Plot', 'Hexbin Plot', 'Box Plot', 'Scatter Plot'], key=f"plot_type_{page}_{i}")
        plot_types.append(plot_type)

        feature1 = st.sidebar.selectbox(f"Select feature for subplot {i+1} on Page {page+1}", df.columns, key=f"feature1_{page}_{i}")
        features1.append(feature1)

        if plot_type in ['Hexbin Plot', 'Scatter Plot']:
            feature2 = st.sidebar.selectbox(f"Select second feature for {plot_type} {i+1} on Page {page+1}", df.columns, key=f"feature2_{page}_{i}")
            features2.append(feature2)
        else:
            features2.append(None)

        title = st.sidebar.text_input(f"Enter title for subplot {i+1} on Page {page+1}", key=f"title_{page}_{i}")
        titles.append(title)

        xlabel = st.sidebar.text_input(f"Enter x-axis label for subplot {i+1} on Page {page+1}", key=f"xlabel_{page}_{i}")
        xlabels.append(xlabel)

        ylabel = st.sidebar.text_input(f"Enter y-axis label for subplot {i+1} on Page {page+1}", key=f"ylabel_{page}_{i}")
        ylabels.append(ylabel)

    page_data.append((plot_types, features1, features2, titles, xlabels, ylabels))

# Create the subplots for each page
figures = []
for page in range(n_pages):
    plot_types, features1, features2, titles, xlabels, ylabels = page_data[page]
    fig = generate_subplots(n_rows, n_cols, df, plot_types, features1, features2, titles, xlabels, ylabels)
    figures.append(fig)

# Export options for the entire dashboard
st.sidebar.header("Export Options")
export_format = st.sidebar.radio("Select export format", ["Word", "PDF"])
if export_format == "Word":
    export_to_word(figures)
elif export_format == "PDF":
    export_to_pdf(figures)
