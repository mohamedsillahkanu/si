import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# App title
st.title("Dummy Table with Categorical and Numeric Variables")

# Sidebar for navigation
st.sidebar.title("Navigation")
section = st.sidebar.radio("Go to", ["Test Overview", "Test Illustration"])

# 1. Test Overview Section
if section == "Test Overview":
    st.header("Overview: Dummy Table")

    st.subheader("What is a Dummy Table?")
    st.write("""
        A dummy table summarizes data involving both categorical and numeric variables.
        It provides insights into how numeric values interact with categorical groups.
    """)

# 2. Test Illustration Section
elif section == "Test Illustration":
    st.header("Illustration: Dummy Table")
    
    st.subheader("Upload your dataset (CSV or XLSX)")
    uploaded_file = st.file_uploader("Choose a CSV or XLSX file", type=["csv", "xlsx"])
    
    if uploaded_file is not None:
        # Load the dataset based on file type
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            elif uploaded_file.name.endswith('.xlsx'):
                df = pd.read_excel(uploaded_file)
                
            st.write("Here is a preview of your data:")
            st.write(df.head())

            # Ask the user to select one categorical and multiple numeric columns
            cat_column = st.selectbox("Select a categorical variable", df.select_dtypes(include=['object', 'category']).columns)
            num_columns = st.multiselect("Select numeric variables", df.select_dtypes(include=['float64', 'int64']).columns)

            # Input field for the title of the dummy table
            table_title = st.text_input("Enter a title for the Dummy Table:", "Summary Table")

            # Generate the summary table
            summary_table = df.groupby(cat_column)[num_columns].agg(['count', 'mean', 'sum', 'std']).reset_index()
            summary_table.columns = [cat_column] + [f"{num_col} ({stat})" for num_col in num_columns for stat in ['Count', 'Mean', 'Total', 'Std Dev']]
            
            # Calculate percentage for each numeric variable
            for num_col in num_columns:
                summary_table[f"{num_col} (Percentage)"] = (summary_table[f"{num_col} (Total)"] / summary_table[f"{num_col} (Total)"].sum()) * 100

            # Display the summary table
            st.write(f"**{table_title}**")
            st.write(summary_table)

            # Input fields for plot titles
            bar_chart_title = st.text_input("Enter a title for the Horizontal Bar Chart:", "Mean by Category")
            pie_chart_title = st.text_input("Enter a title for the Pie Chart:", "Total by Category")

            # Horizontal Bar Charts for Mean
            for num_col in num_columns:
                st.write(f"Horizontal Bar Chart of Mean for {num_col}:")
                mean_values = summary_table.sort_values(by=f"{num_col} (Mean)", ascending=False)
                plt.figure(figsize=(10, 6))
                plt.barh(mean_values[cat_column], mean_values[f"{num_col} (Mean)"], color='skyblue')
                plt.xlabel('Mean')
                plt.title(f'{bar_chart_title} - {num_col}')
                plt.grid(axis='x')
                plt.tight_layout()
                plt.savefig(f'mean_bar_chart_{num_col}.png')
                st.pyplot(plt)

            # Pie Charts for Total
            for num_col in num_columns:
                st.write(f"Pie Chart of Total for {num_col}:")
                plt.figure(figsize=(8, 8))
                plt.pie(mean_values[f"{num_col} (Total)"], labels=mean_values[cat_column], autopct='%1.1f%%', startangle=90)
                plt.title(f'{pie_chart_title} - {num_col}')
                plt.tight_layout()
                plt.savefig(f'total_pie_chart_{num_col}.png')
                st.pyplot(plt)

            # Save all outputs to a Word document
            doc = Document()
            doc.add_heading(table_title, level=1)

            # Add the summary table to the Word document
            doc.add_heading('Summary Table', level=2)
            for i in range(len(summary_table)):
                doc.add_paragraph(str(summary_table.iloc[i].to_dict()))

            # Add plots to the Word document
            for num_col in num_columns:
                doc.add_heading(f'Mean Bar Chart - {num_col}', level=2)
                doc.add_picture(f'mean_bar_chart_{num_col}.png', width=Inches(5))
                
                doc.add_heading(f'Total Pie Chart - {num_col}', level=2)
                doc.add_picture(f'total_pie_chart_{num_col}.png', width=Inches(5))

            # Save the document
            doc.save('dummy_tables.docx')

            # Provide download link for the Word document
            with open('dummy_tables.docx', 'rb') as f:
                st.download_button(
                    label="Download Dummy Tables as Word Document",
                    data=f,
                    file_name='dummy_tables.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

            st.success("Plots have been generated and saved as PNG files, and the Word document has been created.")

        except Exception as e:
            st.error(f"Error loading file: {e}")
