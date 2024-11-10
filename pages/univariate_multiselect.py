import streamlit as st
import seaborn as sns
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Inches

# Function to create a subplot based on user-selected parameters
def create_subplot(ax, data, plot_type, feature1, feature2=None, title="", xlabel="", ylabel="Values"):
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    
    if plot_type == 'Bar Chart':
        if data[feature1].dtype == 'object':  # Ensure it's a categorical variable
            count_data = data[feature1].value_counts()
            sns.barplot(x=count_data.index, y=count_data.values, ax=ax)
        else:
            st.error(f"'{feature1}' is not a categorical variable. Please select a categorical variable for the bar chart.")
    elif plot_type == 'Pie Chart':
        pie_data = data[feature1].value_counts()
        ax.pie(pie_data, labels=pie_data.index, autopct='%1.1f%%')
    elif plot_type == 'Histogram':
        sns.histplot(data[feature1], kde=True, ax=ax)
    elif plot_type == 'Violin Plot':
        sns.violinplot(x=data[feature1], ax=ax)
    elif plot_type == 'Line Plot':
        ax.plot(data[feature1])
    elif plot_type == 'Hexbin Plot' and feature2 is not None:
        ax.hexbin(data[feature1], data[feature2], gridsize=20, cmap='Blues')
    elif plot_type == 'Box Plot':
        sns.boxplot(x=data[feature1], ax=ax)
    elif plot_type == 'Scatter Plot' and feature2 is not None:
        ax.scatter(data[feature1], data[feature2])
        ax.plot([data[feature1].min(), data[feature1].max()], [data[feature1].min(), data[feature1].max()], 'r--')

# Function to generate the subplots based on user inputs
def generate_subplots(rows, cols, data, plot_types, features1, features2, titles, xlabels, ylabels):
    fig, axes = plt.subplots(rows, cols, figsize=(5*cols, 4*rows))

    # If only one subplot, make axes a list
    if rows == 1 and cols == 1:
        axes = [axes]
    elif rows == 1 or cols == 1:
        axes = axes.flatten()  # For single row or single column layouts
    else:
        axes = axes.flatten()  # Flatten axes array for easy iteration

    for i in range(rows * cols):
        if i < len(plot_types):
            plot_type = plot_types[i]
            feature1 = features1[i]
            feature2 = features2[i] if len(features2) > i else None
            title = titles[i] if len(titles) > i else f"Figure {i+1}"
            xlabel = xlabels[i] if len(xlabels) > i else feature1
            ylabel = ylabels[i] if len(ylabels) > i else "Values"
            create_subplot(axes[i], data, plot_type, feature1, feature2, title, xlabel, ylabel)
        else:
            axes[i].axis('off')  # Turn off any unused subplots

    plt.subplots_adjust(wspace=0.4, hspace=0.4)
    st.pyplot(fig)

    return fig  # Return the figure object to export later

# Function to export the dashboard to a Word document
def export_to_word(figures, titles, filename="dashboard.docx"):
    doc = Document()
    doc.add_heading('Dashboard Export', 0)

    for i, fig in enumerate(figures):
        img_stream = BytesIO()
        fig.savefig(img_stream, format='png')
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(6))  # Use Inches from docx.shared
        title = st.text_input(f"Edit title for Figure {i+1}:", titles[i], key=f"title_{i+1}")
        doc.add_paragraph(f'{title}')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(label="Download as Word", data=buffer, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Streamlit app layout
# Initialize df as an empty DataFrame
df = pd.DataFrame()

# Streamlit app layout
st.title("Customizable Data Visualization Dashboard with Editable Titles")

# Sidebar for data source selection
st.sidebar.header("Data Source")
data_source = st.sidebar.radio("Choose data source", ["Use Sample Data", "Upload Excel/CSV"])

if data_source == "Upload Excel/CSV":
    uploaded_file = st.sidebar.file_uploader("Upload an Excel or CSV file", type=["xlsx", "csv"])
    if uploaded_file:
        if uploaded_file.name.endswith(".xlsx"):
            df = pd.read_excel(uploaded_file)
        else:
            df = pd.read_csv(uploaded_file)
        st.write(df)  # Display the dataframe after uploading
else:
    # Sample data
    df = pd.DataFrame({
        'Feature 1': np.random.choice(['A', 'B', 'C', 'D'], 100),
        'Feature 2': np.random.randn(100),
        'Feature 3': np.random.randint(1, 10, 100),
        'Feature 4': np.random.randn(100),
    })

    st.write(df)  # Display sample data

# Sidebar for selecting the chart type and columns for plotting
plot_type = st.sidebar.selectbox("Select plot type", ['Bar Chart', 'Pie Chart', 'Histogram', 'Violin Plot', 'Line Plot', 'Hexbin Plot', 'Box Plot', 'Scatter Plot'])

# MultiSelect for columns
selected_columns = st.sidebar.multiselect("Select columns to plot", df.columns)

# User input for titles and labels
titles = [col for col in selected_columns]
xlabels = [col for col in selected_columns]
ylabels = ["Values" for _ in selected_columns]

# Generate plots for each selected column
figures = []
for col in selected_columns:
    fig = generate_subplots(1, 1, df, [plot_type], [col], [None], [col], [col], ["Values"])
    figures.append(fig)

# Export options
export_option = st.sidebar.selectbox("Export as", ["None", "PDF", "Word"])
if export_option == "Word":
    export_to_word(figures, titles)

