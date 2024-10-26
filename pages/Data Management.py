import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import zipfile
from io import BytesIO
import re
from docx import Document
from docx.shared import Inches

# Sidebar for Data Management
st.sidebar.title("Data Management")
data_management_option = st.sidebar.radio(
    "Choose an option:",
    ["Outlier Detection and Correction"]
)

# Initialize session state variables
if 'df' not in st.session_state:
    st.session_state.df = None
if 'filtered_df' not in st.session_state:
    st.session_state.filtered_df = None

# Function to filter dataframe based on categorical column and unique value
def filter_dataframe(df, categorical_column, unique_value):
    return df[df[categorical_column] == unique_value]

# Function to detect outliers using Scatterplot with Q1 and Q3 lines
def detect_outliers_scatterplot(df, col):
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return df[(df[col] < lower_bound) | (df[col] > upper_bound)]



def replace_outliers_with_mean(df, col):
    # Calculate Q1, Q3, and IQR
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1

    # Define lower and upper bounds for outliers
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR

    # Identify non-outlier values
    non_outliers = df[(df[col] >= lower_bound) & (df[col] <= upper_bound)]

    # Calculate mean of non-outlier values
    mean_value = non_outliers[col].mean()

    # Replace outliers with the calculated mean value
    df[col] = df[col].where((df[col] >= lower_bound) & (df[col] <= upper_bound), mean_value)

    return df





def replace_outliers_with_mean_incld(df, col):
    # Calculate Q1, Q3, and IQR
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1

    # Define lower and upper bounds for outliers
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR


    # Calculate mean of non-outlier values
    mean_value = df[col].mean()

    # Replace outliers with the calculated mean value
    df[col] = df[col].where((df[col] >= lower_bound) & (df[col] <= upper_bound), mean_value)

    return df





def replace_outliers_with_median(df, col):
    # Calculate Q1, Q3, and IQR
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1

    # Define lower and upper bounds for outliers
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR

    # Identify non-outlier values
    non_outliers = df[(df[col] >= lower_bound) & (df[col] <= upper_bound)]

    # Calculate mean of non-outlier values
    mean_value = non_outliers[col].median()

    # Replace outliers with the calculated mean value
    df[col] = df[col].where((df[col] >= lower_bound) & (df[col] <= upper_bound), mean_value)

    return df





def replace_outliers_with_median_incld(df, col):
    # Calculate Q1, Q3, and IQR
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1

    # Define lower and upper bounds for outliers
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR


    # Calculate mean of non-outlier values
    mean_value = df[col].median()

    # Replace outliers with the calculated mean value
    df[col] = df[col].where((df[col] >= lower_bound) & (df[col] <= upper_bound), mean_value)

    return df



# Function to apply Winsorization to a column
def winsorize_column(df, col):
    # Calculate the lower and upper percentiles
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1

    # Define lower and upper bounds for outliers
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR

    # Winsorize the column
    df[col] = df[col].clip(lower=lower_bound, upper=upper_bound)

    return df


# Function to calculate moving average excluding outliers
def calculate_moving_avg_excluding_outliers(series, window, lower_bound, upper_bound):
    return series.rolling(window=window, min_periods=1, center=True).apply(
        lambda x: np.mean(x[(x >= lower_bound) & (x <= upper_bound)]) if not x[(x >= lower_bound) & (x <= upper_bound)].empty else np.nan
    )

# Function to apply moving average and correct outliers
def replace_outliers_with_corrected_moving_avg(df, col):
    # Calculate Q1, Q3, and IQR for detecting outliers
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1

    # Define lower and upper bounds for outliers
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR

    # Set the window size for moving average
    window_size = 3

    # Calculate the moving average excluding outliers
    df[f'{col}_Moving_Avg_Excluded_Outliers'] = calculate_moving_avg_excluding_outliers(df[col], window_size, lower_bound, upper_bound)

    # Create a corrected column replacing outliers with the calculated moving average
    df[f'{col}_corrected_Moving_Avg_Excluded_Outliers'] = np.where(
        (df[col] < lower_bound) | (df[col] > upper_bound),
        df[f'{col}_Moving_Avg_Excluded_Outliers'],
        df[col]
    )

    return df




  # Function to calculate moving average including all values
def calculate_moving_avg_including_outliers(series, window):
    return series.rolling(window=window, min_periods=1, center=True).mean()

# Function to apply moving average and correct outliers
def replace_outliers_with_corrected_moving_avg_including(df, col):
    # Calculate Q1, Q3, and IQR for detecting outliers
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1

    # Define lower and upper bounds for outliers
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR

    # Set the window size for moving average
    window_size = 3

    # Calculate the moving average including all values
    df[f'{col}_Moving_Avg_Including_Outliers'] = calculate_moving_avg_including_outliers(df[col], window_size)

    # Create a corrected column replacing outliers with the calculated moving average
    df[f'{col}_corrected_Moving_Avg_Including_Outliers'] = np.where(
        (df[col] < lower_bound) | (df[col] > upper_bound),
        df[f'{col}_Moving_Avg_Including_Outliers'],
        df[col]
    )

    return df





# Function to clean sheet names (remove invalid characters)
def clean_sheet_name(sheet_name):
    # Remove invalid characters for Excel sheet names
    clean_name = re.sub(r'[\\/*?[\]:]', '_', sheet_name)
    return clean_name[:31]  # Excel sheet names must be <= 31 characters

# Main App
st.title("Outlier Detection and Correction")

if data_management_option == "Outlier Detection and Correction":
    st.header("Select Columns for Outlier Detection")

    if st.session_state.df is None:
        st.warning("No dataset available. Please upload your dataset.")
        uploaded_file = st.file_uploader("Upload a CSV or Excel file", type=["csv", "xlsx"])
        if uploaded_file:
            if uploaded_file.name.endswith(".csv"):
                st.session_state.df = pd.read_csv(uploaded_file)
            else:
                st.session_state.df = pd.read_excel(uploaded_file)
            st.success("Dataset uploaded successfully!")
            st.dataframe(st.session_state.df)

    if st.session_state.df is not None:
        # User selects a categorical and numerical column
        categorical_column = st.selectbox("Select a Categorical Column:", st.session_state.df.select_dtypes(include='object').columns)
        numeric_column = st.selectbox("Select a Numeric Column for Outlier Detection:", st.session_state.df.select_dtypes(include='number').columns)

        if categorical_column and numeric_column:
            st.write("Outliers will be detected and visualized for each unique value in the categorical column.")

            # Filter data for each unique value in the categorical column
            unique_values = st.session_state.df[categorical_column].unique()
            word_buffer = BytesIO()  # Create a buffer for the Word file
            excel_buffer = BytesIO()  # Create a buffer for the Excel file
            document = Document()

            outlier_data = pd.ExcelWriter(excel_buffer, engine='xlsxwriter')

            # Prepare the zip archive for the output files
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zf:

                # Loop through each unique value
                for i, value in enumerate(unique_values):
                    filtered_df = filter_dataframe(st.session_state.df, categorical_column, value)

                    # Detect outliers in the numeric column using scatter plot method
                    outliers = detect_outliers_scatterplot(filtered_df, numeric_column)

                    # Ensure outliers and only necessary columns are included in Excel
                    if not outliers.empty:
                        outlier_subset = outliers[[categorical_column, numeric_column]]
                        # Clean the sheet name and ensure it's <= 31 characters
                        clean_value = clean_sheet_name(str(value))
                        # Write the outliers to the Excel file
                        outlier_subset.to_excel(outlier_data, sheet_name=clean_value, index=False)

                    # Create scatter plot with Q1 and Q3 lines for the Word document
                    fig, ax = plt.subplots(figsize=(8, 6))
                    ax.scatter(filtered_df.index, filtered_df[numeric_column], label='Data Points')
                    Q1 = filtered_df[numeric_column].quantile(0.25)
                    Q3 = filtered_df[numeric_column].quantile(0.75)
                    IQR = Q3 - Q1
                    lower_bound = Q1 - 1.5 * IQR
                    upper_bound = Q3 + 1.5 * IQR

                    ax.axhline(lower_bound, color='blue', linestyle='--', label='Lower Bound')
                    ax.axhline(upper_bound, color='red', linestyle='--', label='Upper Bound')
                    ax.set_title(f'Scatterplot for {numeric_column} (Category: {value})')
                    ax.set_xlabel('Index')
                    ax.set_ylabel(numeric_column)
                    ax.legend()

                    # Save plot as image to add to Word document
                    img_buffer = BytesIO()
                    plt.savefig(img_buffer, format='png')
                    img_buffer.seek(0)
                    document.add_paragraph(f'Outlier Detection for {categorical_column} = {value}')
                    document.add_picture(img_buffer, width=Inches(5))
                    plt.close(fig)

                # Save the Word document
                document.save(word_buffer)

                # Close the Excel writer (finalize the file)
                outlier_data.close()

                # Write Word and Excel to zip
                zf.writestr('outliers_report.docx', word_buffer.getvalue())
                zf.writestr('outliers_data.xlsx', excel_buffer.getvalue())

            st.success("Outliers processed successfully!")
            # Offer the zip file for download
            st.download_button(
                label="Download ZIP file with Word Report and Outliers",
                data=zip_buffer.getvalue(),
                file_name='outliers_results.zip',
                mime='application/zip'
            )
